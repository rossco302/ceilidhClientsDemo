from kivy.app import App
from datetime import date
from docx import Document
from docx2pdf import convert
from kivy.uix.textinput import TextInput
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.properties import ObjectProperty
from kivy.clock import Clock
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.recycleview import RecycleView
from kivy.uix.behaviors import FocusBehavior
from kivy.uix.recycleview.layout import LayoutSelectionBehavior
from kivy.uix.recycleboxlayout import RecycleBoxLayout
from kivy.properties import ListProperty
from kivy.uix.popup import Popup
from kivy.storage.jsonstore import JsonStore
from kivy.core.window import Window
from kivy.core.window import Window


clients_store = JsonStore('clients.json')
Window.size = (1100, 650)



def populate_rv(dt):
	app = App.get_running_app()
	app.clients = [{'text': str(clients_store.get(x)['client_id']) + f' - {str(x)}'} for x in clients_store if x != 'client_counter']

Clock.schedule_once(populate_rv, 0.5)

clients = []

class CeilidhClientsDemoApp(App):
	clients = ListProperty()
	#Window.clearcolor = (1, 1, 1, 1)


def convert_call_back(dt):
	pass
	#convert('clydeside_celidh_contract_draft_gui_completed.docx', 'clydeside_celidh_contract_draft_gui_completed.pdf')

class ScreenManager(ScreenManager):
	pass

class HomeScreen(Screen):
	'''should have a rv containg the clients
	should have a pop up to add client name, this should also add a blank entry into the data base
	should have the ability to click on each client and open the form page with the details filled in from the clients section of the data base'''
	def populate_rv(dt):
		app = App.get_running_app()
		app.clients = [{'text': str(clients_store.get(x)['client_id']) + f' - {str(x)}'} for x in clients_store if x != 'client_counter']

	def on_pre_enter(self, *largs):
		app = App.get_running_app()
		app.clients = [{'text': str(clients_store.get(x)['client_id']) + f' - {str(x)}'} for x in clients_store if x != 'client_counter']

class FormScreen(Screen):

	def create_deposit_invoice(self):
		document = Document('word_drafts/deposit_invoice_draft.docx')

		client_name = self.ids.client_name.text
		travel_info = self.ids.travel_info.text
		services = self.ids.services.text
		ceilidh_date = self.ids.ceilidh_date.text
		timings = self.ids.timings.text
		venue = self.ids.venue.text
		deposit = (int(self.ids.fee_per_musician.text) * int(self.ids.num_of_musicians.text) + int(self.ids.booking_fee.text)) * 0.25
		print('deposit:', f'£{deposit}')

		date_today = date.today()
		document.paragraphs[6].runs[3].text = date_today.strftime('%d/%m/%y')
		document.paragraphs[9].runs[2].text = client_name
		document.paragraphs[14].runs[3].text = services
		document.paragraphs[12].runs[5].text = ceilidh_date
		document.paragraphs[13].runs[3].text = timings
		document.paragraphs[15].runs[3].text = venue
		document.paragraphs[7].runs[2].text = str(clients_store.get(client_name)['client_id'])

		try:
			document.paragraphs[16].runs[1].text = f"£{str(deposit)}"
		except:
			document.paragraphs[16].runs[1].text = ''


		document.save(f"clydeside_celidh_deposit_invoice_{client_name.replace(' ', '_').lower()}.docx")


	def create_invoice(self):
		document = Document('word_drafts/invoice_draft.docx')

		client_name = self.ids.client_name.text
		travel_info = self.ids.travel_info.text
		services = self.ids.services.text
		ceilidh_date = self.ids.ceilidh_date.text
		timings = self.ids.timings.text
		venue = self.ids.venue.text
		musicians_fee = int(self.ids.num_of_musicians.text) * int(self.ids.fee_per_musician.text)
		total_fee = musicians_fee + int(self.ids.booking_fee.text)
		deposit = total_fee * 0.25

		try:
			remaining_fee = total_fee + int(self.ids.travel_costs.text) + int(self.ids.accomodation_cost.text) - deposit
		except:
			remaining_fee = total_fee + int(self.ids.accomodation_cost.text) - deposit

		print(remaining_fee)

		date_today = date.today()
		document.paragraphs[7].runs[3].text = date_today.strftime('%d/%m/%y')
		document.paragraphs[10].runs[3].text = client_name
		document.paragraphs[15].runs[2].text = services
		document.paragraphs[13].runs[4].text = ceilidh_date
		document.paragraphs[14].runs[2].text = timings
		document.paragraphs[16].runs[2].text = venue
		document.paragraphs[17].runs[1].text = str(remaining_fee)
		document.paragraphs[8].runs[1].text = str(clients_store.get(client_name)['client_id'])

		document.save(f"clydeside_celidh_invoice_{client_name.replace(' ', '_').lower()}.docx")

	def delete_client(self):
		clients_store.delete(client_pressed)
		app = App.get_running_app()
		app.clients = [{'text': str(x)} for x in clients_store]
		clients_store.get('client_counter')['client_id'] -= 1

	def on_pre_enter(self, *largs):
		app = App.get_running_app()
		self.ids.client_name.text = client_pressed
		self.ids.services.text = clients_store.get(client_pressed)['services']
		self.ids.ceilidh_date.text = clients_store.get(client_pressed)['ceilidh_date']
		self.ids.timings.text = clients_store.get(client_pressed)['start_time']
		self.ids.accomodation.text = clients_store.get(client_pressed)['accomodation']
		self.ids.venue.text = clients_store.get(client_pressed)['venue']
		self.ids.fee_per_musician.text = clients_store.get(client_pressed)['fee_per_musician']
		self.ids.accomodation_cost.text = clients_store.get(client_pressed)['accomodation_cost']
		self.ids.travel_info.text = clients_store.get(client_pressed)['travel_info']
		self.ids.num_of_musicians.text = clients_store.get(client_pressed)['num_of_musicians']
		self.ids.booking_fee.text = clients_store.get(client_pressed)['booking_fee']
		self.ids.travel_costs.text = clients_store.get(client_pressed)['travel_costs']


	def save_clients_info(self):
		clients_store.put(self.ids.client_name.text, travel_costs = self.ids.travel_costs.text, booking_fee = self.ids.booking_fee.text, num_of_musicians = self.ids.num_of_musicians.text, client_id = clients_store.get(self.ids.client_name.text)['client_id'], services = self.ids.services.text, ceilidh_date = self.ids.ceilidh_date.text, start_time = self.ids.timings.text, accomodation = self.ids.accomodation.text, venue = self.ids.venue.text, fee_per_musician = self.ids.fee_per_musician.text, accomodation_cost = self.ids.accomodation_cost.text, travel_info = self.ids.travel_info.text)

	def create_contract(self):
		total_fee = ObjectProperty()
		'''
		set the document variable to the draft document
		get the desired locations of the form input.
		set the text attribute of each text box to the variable names
		calculate the variable that re not inputed (such as deposit and total fee)
		set the text attribute of these locations to the text attributes of the text inputs
		save the document
		create a pdf version
		'''

		#set the document variable to the draft document
		document = Document('word_drafts/contract_draft.docx')

		#get the locations in the word draft
		date_today_correct_format_run = document.paragraphs[3].runs[2]
		client_name_parathree_run = document.paragraphs[3].runs[5]
		client_name_paraseven_run = document.paragraphs[7].runs[2]
		services_run = document.paragraphs[8].runs[1]
		ceilidh_date_run = document.paragraphs[9].runs[1]
		start_time_run = document.paragraphs[10].runs[2]
		venue_para_eleven_run = document.paragraphs[11].runs[1]
		venue_para_four_run = document.paragraphs[4].runs[5]
		accomodation_run = document.paragraphs[12].runs[2]
		musicians_fee_run = document.paragraphs[15].runs[5]
		accomodation_cost_run = document.paragraphs[18].runs[4]
		total_fee_run = document.paragraphs[19].runs[4]
		deposit_run = document.paragraphs[21].runs[7]
		travel_info_run = document.paragraphs[17].runs[8]
		booking_fee_run = document.paragraphs[16].runs[5]

		#set the inputed variables
		client_name = self.ids.client_name.text
		services = self.ids.services.text
		ceilidh_date = self.ids.ceilidh_date.text
		timings = self.ids.timings.text
		accomodation = self.ids.accomodation.text
		venue = self.ids.venue.text
		fee_per_musician = self.ids.fee_per_musician.text
		num_of_musicians = self.ids.num_of_musicians.text
		accomodation_cost = self.ids.accomodation_cost.text
		travel_cost = self.ids.travel_costs.text
		booking_fee = self.ids.booking_fee.text

		#calculate the non inputed variables
		date_today = date.today()
		date_today_correct_format = date_today.strftime('%d/%m/%y')

		try: 
			int(travel_cost)
			total_fee = (int(fee_per_musician) * int(num_of_musicians)) + int(self.ids.booking_fee.text) + int(travel_cost)
		except:
			total_fee = (int(fee_per_musician) * int(num_of_musicians)) + int(self.ids.booking_fee.text)

		print(f'total fee {total_fee}')

		deposit = int(int(total_fee) * 0.25)

		#set the text attributes in the word doc
		date_today_correct_format_run.text = date_today_correct_format
		client_name_parathree_run.text = client_name
		client_name_paraseven_run.text = client_name
		services_run.text = services
		ceilidh_date_run.text = ceilidh_date
		start_time_run.text = timings
		venue_para_eleven_run.text = venue
		venue_para_four_run.text = venue
		accomodation_run.text = accomodation
		musicians_fee_run.text = str(int(fee_per_musician) * int(num_of_musicians)) 
		accomodation_cost_run.text = str(accomodation_cost)
		total_fee_run.text = str(total_fee)
		deposit_run.text = str(deposit)
		travel_info_run.text = travel_cost
		booking_fee_run.text = booking_fee

		#save the document
		document.save(f"clydeside_celidh_contract_{client_pressed.replace(' ', '_').lower()}.docx")

		#convert to pdf
		Clock.schedule_once(convert_call_back, 1)

class ClientsRv(RecycleView):
	def __init__(self, **kwargs):
		super(ClientsRv, self).__init__(**kwargs)

class SelectableRecycleBoxLayout(FocusBehavior, LayoutSelectionBehavior, RecycleBoxLayout):
	pass

class AddClientPopUp(Popup):
	new_client = ObjectProperty()
	
	def add_new_client(self):
		#add the new name into the data bass, then update the rv data from the data bass
		clients_store.get('client_counter')['client_id'] += 1
		clients_store.put(self.ids.new_client.text, travel_costs = '', booking_fee = '', num_of_musicians = '', client_id = clients_store.get('client_counter')['client_id'], services = '', ceilidh_date = '', start_time = '', accomodation = '', venue = '', fee_per_musician = '', accomodation_cost = '', travel_info = '')
		app = App.get_running_app()
		app.clients = [{'text': str(clients_store.get(x)['client_id']) + f' - {str(x)}'} for x in clients_store if x != 'client_counter']


class ClientsRvButton(Button):
	def get_client_data(self):
		global client_pressed 
		client_pressed = self.text.split('-')[1].lstrip()


class CeildhInfoForm(BoxLayout):
	client_name = ObjectProperty()
	services = ObjectProperty()
	ceilidh_date = ObjectProperty()
	timings = ObjectProperty()
	accomodation = ObjectProperty()
	venue = ObjectProperty()
	musicians_fee = ObjectProperty()
	accomodation_cost = ObjectProperty()
	travel_info = ObjectProperty()
	num_of_musicians = ObjectProperty()
	booking_fee = ObjectProperty()
	travel_costs = ObjectProperty()

if __name__ == '__main__':
	CeilidhClientsDemoApp().run()

